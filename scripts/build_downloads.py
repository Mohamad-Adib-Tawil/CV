from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "assets" / "cv"
OUTPUT_DIR = ROOT / "assets" / "downloads"

FILES = {
    "CV_EN.docx": ("CV_Flutter_EN.txt", False),
    "CV_AR.docx": ("CV_Flutter_AR.txt", True),
    "CV_SE_EN.docx": ("CV_SE_EN.txt", False),
    "CV_SE_AR.docx": ("CV_SE_AR.txt", True),
}

ACCENT = RGBColor(30, 64, 175)
MUTED = RGBColor(71, 85, 105)
BODY = RGBColor(15, 23, 42)
DIVIDER_RE = re.compile(r"^[=-]{20,}$")
URL_RE = re.compile(r"https?://\S+")
LABEL_RE = re.compile(r"^[^:]{1,28}:\s+")


def set_run_font(run, *, size: float = 11, bold: bool = False, color=BODY) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0)
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    for run in paragraph.runs:
        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        run._element.get_or_add_rPr().append(rtl)


def add_hyperlinked_text(paragraph, text: str, is_arabic: bool) -> None:
    cursor = 0
    for match in URL_RE.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor : match.start()]))

        url = match.group(0).rstrip(".,)")
        relationship = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship)
        hyperlink_run = OxmlElement("w:r")
        run_props = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1D4ED8")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_props.extend([color, underline])
        hyperlink_run.append(run_props)
        text_node = OxmlElement("w:t")
        text_node.text = url
        hyperlink_run.append(text_node)
        hyperlink.append(hyperlink_run)
        paragraph._p.append(hyperlink)
        cursor = match.start() + len(url)

    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]))
    if is_arabic:
        set_rtl(paragraph)


def configure_document(document: Document, is_arabic: bool) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in (("Heading 1", 13), ("Heading 2", 11.5), ("Heading 3", 10.5)):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(10)
    bullet.paragraph_format.left_indent = Inches(0.34)
    bullet.paragraph_format.first_line_indent = Inches(-0.18)
    bullet.paragraph_format.space_after = Pt(2.5)
    bullet.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("Mohamad Adib Tawil  •  CV"), size=8.5, color=MUTED)
    if is_arabic:
        set_rtl(footer)


def normalize_lines(source: Path) -> list[str]:
    raw = source.read_text(encoding="utf-8").splitlines()
    return [line.rstrip() for line in raw if not DIVIDER_RE.match(line.strip())]


def is_section_heading(line: str, previous_blank: bool, next_blank: bool) -> bool:
    stripped = line.strip()
    if not previous_blank or not next_blank or len(stripped) > 64 or ":" in stripped:
        return False
    if re.search(r"[A-Za-z]", stripped):
        return stripped == stripped.upper()
    return not stripped.startswith("-") and len(stripped.split()) <= 6


def build_document(source_name: str, output_name: str, is_arabic: bool) -> None:
    lines = normalize_lines(SOURCE_DIR / source_name)
    document = Document()
    configure_document(document, is_arabic)

    content = [line for line in lines if line.strip()]
    title = content[0]
    subtitle = content[1]

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_arabic else WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.space_after = Pt(2)
    set_run_font(title_paragraph.add_run(title), size=23, bold=True, color=ACCENT)
    if is_arabic:
        set_rtl(title_paragraph)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = title_paragraph.alignment
    subtitle_paragraph.paragraph_format.space_after = Pt(10)
    set_run_font(subtitle_paragraph.add_run(subtitle), size=12.5, bold=True, color=MUTED)
    if is_arabic:
        set_rtl(subtitle_paragraph)

    start_index = lines.index(content[1]) + 1
    index = start_index
    previous_blank = True

    while index < len(lines):
        line = lines[index].strip()
        next_blank = index + 1 >= len(lines) or not lines[index + 1].strip()
        if not line:
            previous_blank = True
            index += 1
            continue

        if is_section_heading(line, previous_blank, next_blank):
            paragraph = document.add_paragraph(style="Heading 1")
            set_run_font(paragraph.add_run(line), size=13, bold=True, color=ACCENT)
            if is_arabic:
                set_rtl(paragraph)
            previous_blank = False
            index += 1
            continue

        if line.startswith("- "):
            parts = [line[2:].strip()]
            cursor = index + 1
            while cursor < len(lines):
                continuation = lines[cursor]
                if not continuation.strip() or continuation.lstrip().startswith("-"):
                    break
                if is_section_heading(
                    continuation.strip(),
                    previous_blank=False,
                    next_blank=cursor + 1 >= len(lines) or not lines[cursor + 1].strip(),
                ):
                    break
                parts.append(continuation.strip())
                cursor += 1
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.keep_together = True
            add_hyperlinked_text(paragraph, " ".join(parts), is_arabic)
            index = cursor
            previous_blank = False
            continue

        parts = [line]
        cursor = index + 1
        while cursor < len(lines) and lines[cursor].strip():
            candidate = lines[cursor].strip()
            if LABEL_RE.match(candidate) or candidate.startswith("- ") or is_section_heading(
                candidate,
                previous_blank=False,
                next_blank=cursor + 1 >= len(lines) or not lines[cursor + 1].strip(),
            ):
                break
            parts.append(candidate)
            cursor += 1

        text = " ".join(parts)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = False
        if ":" in text and text.index(":") < 28:
            label, value = text.split(":", 1)
            set_run_font(paragraph.add_run(f"{label}:"), bold=True)
            add_hyperlinked_text(paragraph, value, is_arabic)
        else:
            add_hyperlinked_text(paragraph, text, is_arabic)
        index = cursor
        previous_blank = False

    core_props = document.core_properties
    core_props.title = f"{title} - {subtitle}"
    core_props.author = "Mohamad Adib Tawil"
    core_props.subject = "Professional CV"
    core_props.keywords = "Flutter, Dart, Software Engineer, Mobile Developer"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DIR / output_name)


def main() -> None:
    for output_name, (source_name, is_arabic) in FILES.items():
        build_document(source_name, output_name, is_arabic)
        print(f"Generated {OUTPUT_DIR / output_name}")


if __name__ == "__main__":
    main()
